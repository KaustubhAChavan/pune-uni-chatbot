import os
import PyPDF2
from typing import List, Dict, Any
from pathlib import Path
import docx
import re

# Import necessary LangChain components
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.docstore.document import Document

class KnowledgeBase:
    def __init__(self, data_dir: Path, vector_store_dir: Path, openai_api_key: str):
        """Initialize the knowledge base with data directory and vector store directory."""
        self.data_dir = data_dir
        self.vector_store_dir = vector_store_dir
        self.embeddings = OpenAIEmbeddings(openai_api_key=openai_api_key)
        self.vector_store = None
        
    def load_documents(self) -> List[Document]:
        """Load documents from data directory."""
        documents = []
        
        # Walk through the data directory
        for root, _, files in os.walk(self.data_dir):
            for file in files:
                file_path = os.path.join(root, file)
                
                # Process PDF files
                if file.lower().endswith('.pdf'):
                    text = self._extract_text_from_pdf(file_path)
                    if text:
                        documents.append(Document(page_content=text, metadata={"source": file_path}))
                
                # Process DOCX files
                elif file.lower().endswith('.docx'):
                    text = self._extract_text_from_docx(file_path)
                    if text:
                        documents.append(Document(page_content=text, metadata={"source": file_path}))
                
                # Process TXT files
                elif file.lower().endswith('.txt'):
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        text = f.read()
                    if text:
                        documents.append(Document(page_content=text, metadata={"source": file_path}))
        
        return documents
    
    def _extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from a PDF file."""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"
            return text
        except Exception as e:
            print(f"Error extracting text from PDF {file_path}: {str(e)}")
            return ""
    
    def _extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from a DOCX file."""
        try:
            doc = docx.Document(file_path)
            full_text = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                full_text.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text.append(' | '.join(row_text))
            
            return '\n'.join(full_text)
        except Exception as e:
            print(f"Error extracting text from DOCX {file_path}: {str(e)}")
            return ""
    
    def create_or_load_vector_store(self, force_reload: bool = False) -> FAISS:
        """Create or load the vector store."""
        vector_store_path = self.vector_store_dir / "pune_university_faiss"
        
        # If vector store exists and not forced to reload, load it
        if vector_store_path.exists() and not force_reload:
            try:
                self.vector_store = FAISS.load_local(
                    str(vector_store_path), 
                    self.embeddings,
                    allow_dangerous_deserialization=True  # Add this line
                )
                print("Loaded existing vector store.")
                return self.vector_store
            except Exception as e:
                print(f"Error loading vector store: {str(e)}. Creating new one.")
        
        # Create new vector store
        print("Creating new vector store...")
        documents = self.load_documents()
        print(f"Loaded {len(documents)} documents.")
        
        if not documents:
            raise ValueError("No documents found in the data directory.")
        
        # Split documents into chunks
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
        )
        chunks = text_splitter.split_documents(documents)
        print(f"Split into {len(chunks)} chunks.")
        
        # Create vector store
        self.vector_store = FAISS.from_documents(chunks, self.embeddings)
        
        # Save vector store
        self.vector_store.save_local(str(vector_store_path))
        print(f"Vector store saved to {vector_store_path}.")
        
        return self.vector_store
    
    def query_knowledge_base(self, query: str, k: int = 5) -> List[Dict[str, Any]]:
        """Query the knowledge base for relevant information."""
        if not self.vector_store:
            self.create_or_load_vector_store()
        
        # Perform similarity search
        docs_and_scores = self.vector_store.similarity_search_with_score(query, k=k)
        
        results = []
        for doc, score in docs_and_scores:
            # Extract filename from path
            source = doc.metadata.get("source", "Unknown")
            filename = os.path.basename(source) if source != "Unknown" else "Unknown"
            
            # Add to results
            results.append({
                "content": doc.page_content,
                "source": filename,
                "score": float(score)  # Convert to float for JSON serialization
            })
        
        return results